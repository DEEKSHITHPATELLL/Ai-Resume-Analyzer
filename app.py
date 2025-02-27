import os
import io
import re
import json
import logging
import tempfile
from flask import Flask, request, jsonify, render_template
from werkzeug.utils import secure_filename
import PyPDF2
import docx2txt
from docx import Document
from dotenv import load_dotenv
import google.generativeai as genai
import numpy as np
import pandas as pd
import shap

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Load environment variables
load_dotenv()

# Configure Google Generative AI
GOOGLE_API_KEY = os.getenv('GOOGLE_API_KEY')
if not GOOGLE_API_KEY:
    raise ValueError("GOOGLE_API_KEY not found in environment variables")

app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024 
app.config['UPLOAD_FOLDER'] = 'uploads'

ALLOWED_EXTENSIONS = {'pdf', 'docx'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def extract_text_from_pdf(file_content):
    """Extract text from PDF using multiple methods including OCR"""
    try:
        logger.info(f"Starting PDF extraction, content size: {len(file_content)} bytes")
        
        # Try extracting with PyPDF2 first
        try:
            reader = PyPDF2.PdfReader(io.BytesIO(file_content))
            logger.info(f"PDF has {len(reader.pages)} pages")
            
            if reader.is_encrypted:
                logger.error("PDF is password protected")
                return None
                
            text_parts = []
            for i, page in enumerate(reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text and page_text.strip():
                        text_parts.append(page_text.strip())
                        logger.info(f"Page {i+1}: Extracted {len(page_text)} characters")
                    else:
                        logger.warning(f"Page {i+1}: No text extracted, might be scanned")
                except Exception as e:
                    logger.error(f"Error on page {i+1}: {str(e)}")
                    continue
            
            # If we got meaningful text, return it
            if text_parts:
                text = '\n'.join(text_parts)
                if len(text.strip()) > 50:
                    logger.info(f"Successfully extracted {len(text)} characters with PyPDF2")
                    return text
                    
            logger.info("PyPDF2 extraction yielded no meaningful text, trying OCR...")
            
        except Exception as e:
            logger.error(f"PyPDF2 extraction failed: {str(e)}")
            
        # Try OCR as fallback
        try:
            import pytesseract
            import fitz  # PyMuPDF
            from PIL import Image
            import os
            
            # Check for Tesseract installation
            tesseract_path = "C:\\Program Files\\Tesseract-OCR\\tesseract.exe"
            if not os.path.exists(tesseract_path):
                logger.error(f"Tesseract not found at {tesseract_path}")
                logger.error("Please install from: https://github.com/UB-Mannheim/tesseract/wiki")
                return None
                
            # Set Tesseract path
            pytesseract.pytesseract.tesseract_cmd = tesseract_path
            
            # Try OCR using PyMuPDF
            try:
                # Load PDF
                pdf_document = fitz.open(stream=file_content, filetype="pdf")
                
                # Get first page
                page = pdf_document[0]
                
                # Convert to image
                pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))  # 2x zoom for better quality
                
                # Convert to PIL Image
                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                
                # Convert to grayscale for better OCR
                img = img.convert('L')
                
                # Perform OCR
                text = pytesseract.image_to_string(img)
                
                if text and len(text.strip()) > 50:
                    logger.info(f"Successfully extracted {len(text)} characters with OCR")
                    return text.strip()
                    
                logger.warning("OCR yielded no meaningful text")
                
            except Exception as e:
                logger.error(f"Error during OCR: {str(e)}")
            
            logger.error("OCR extraction failed")
            return None
            
        except ImportError as e:
            logger.error(f"OCR dependencies not installed: {str(e)}")
            return None
        except Exception as e:
            logger.error(f"OCR extraction failed: {str(e)}")
            return None
            
    except Exception as e:
        logger.error(f"PDF extraction failed: {str(e)}")
        return None

def extract_text_from_docx(file_content):
    """Extract text from DOCX using multiple methods"""
    try:
        logger.info(f"Starting DOCX extraction, content size: {len(file_content)} bytes")

        # Create temporary file
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_file:
            temp_file.write(file_content)
            temp_file.flush()
            temp_path = temp_file.name
            
        try:
            # Try docx2txt first
            try:
                text = docx2txt.process(temp_path)
                if text and len(text.strip()) > 50:
                    logger.info(f"Successfully extracted {len(text)} characters with docx2txt")
                    return text.strip()
                else:
                    logger.warning("docx2txt extraction yielded no meaningful text")
            except Exception as e:
                logger.error(f"docx2txt extraction failed: {str(e)}")
            
            # Try python-docx as fallback
            doc = Document(temp_path)
            text_parts = []
            
            # Extract from paragraphs
            for para in doc.paragraphs:
                if para.text and para.text.strip():
                    text_parts.append(para.text.strip())
                    
            # Extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text and cell.text.strip():
                            text_parts.append(cell.text.strip())
            
            if text_parts:
                text = '\n'.join(text_parts)
                if len(text.strip()) > 50:
                    logger.info(f"Successfully extracted {len(text)} characters with python-docx")
                    return text
                    
            logger.error("No meaningful text extracted from DOCX")
            return None
            
        finally:
            try:
                os.unlink(temp_path)
            except Exception as e:
                logger.error(f"Error deleting temp file: {str(e)}")
                
    except Exception as e:
        logger.error(f"DOCX extraction failed: {str(e)}")
        return None

def analyze_resume_with_ai(text):
    """Analyze resume text using Google's Generative AI"""
    try:
        model = genai.GenerativeModel('gemini-1.5-pro')
        skills_prompt = f"""Analyze this resume and extract information. 
Respond ONLY with a JSON object in this EXACT format:
{{
    "skills": ["skill1", "skill2", "skill3"],
    "experience": "Brief summary of work experience",
    "education": "Brief summary of education"
}}

Note: For skills, extract ALL technical skills, soft skills, and domain knowledge.
Include programming languages, frameworks, tools, methodologies, and soft skills.

Resume text:
{text}"""

        skills_response = model.generate_content(skills_prompt)
        logger.info(f"AI Skills Response: {skills_response.text}")
        
        # Clean the response text to ensure it's valid JSON
        response_text = skills_response.text.strip()
        if response_text.startswith('```json'):
            response_text = response_text[7:]
        if response_text.endswith('```'):
            response_text = response_text[:-3]
        response_text = response_text.strip()
        
        try:
            extracted_info = json.loads(response_text)
            if not isinstance(extracted_info.get('skills'), list):
                raise ValueError("Skills must be a list")
            logger.info(f"Successfully parsed skills: {extracted_info['skills']}")
        except Exception as e:
            logger.error(f"Failed to parse skills JSON: {str(e)}")
            extracted_info = {
                "skills": [],
                "experience": "Failed to extract experience",
                "education": "Failed to extract education"
            }
        guidance_prompt = f"""Based on the candidate's skills, experience, and education, provide detailed career guidance.
Focus on creating UNIQUE and PERSONALIZED recommendations. DO NOT use generic placeholders.

For each career path and job opportunity, ensure they are DIFFERENT and SPECIFIC to the candidate's background.

The response must be a JSON object following this EXACT format:
{{
    "career_paths": [
        {{
            "role": "Primary Career Path - Most Suitable Role",
            "description": "Detailed explanation of why this is the best path based on their current skills",
            "growth_potential": "Specific growth trajectory with timeline and positions",
            "salary_range": "Specific salary range with location context",
            "next_steps": [
                "Immediate action item to start this path",
                "Short-term goal (3-6 months)",
                "Medium-term goal (6-12 months)",
                "Long-term goal (1-2 years)"
            ]
        }},
        {{
            "role": "Alternative Career Path - Different Industry/Role",
            "description": "Why this alternative path could be interesting based on transferable skills",
            "growth_potential": "Different growth trajectory in this alternative path",
            "salary_range": "Salary range in this alternative path",
            "next_steps": [
                "First step to transition to this path",
                "Key skills to develop",
                "Industry connections to make",
                "Timeline for transition"
            ]
        }}
    ],
    "recommendations": [
        "Specific immediate certification or course to pursue",
        "Concrete project to build for portfolio",
        "Industry-specific networking suggestion",
        "Technical skill to prioritize learning"
    ],
    "job_opportunities": [
        {{
            "title": "Immediate Job Opportunity",
            "description": "Role matching current skills - can apply immediately",
            "required_skills": ["Current technical skill", "Existing soft skill", "Required domain knowledge"],
            "matching_skills": ["Skills from their profile that match requirements"],
            "missing_skills": ["Minor skills to develop while working"],
            "typical_salary": "Current market salary range",
            "demand_level": "Current market demand with growth rate"
        }},
        {{
            "title": "Future Job Opportunity - After Upskilling",
            "description": "Role to target after gaining additional skills",
            "required_skills": ["Advanced technical skill", "Leadership skill", "Specialized domain knowledge"],
            "matching_skills": ["Transferable skills they already have"],
            "missing_skills": ["Key skills to develop for this role"],
            "typical_salary": "Higher salary range after upskilling",
            "demand_level": "Future market demand projection"
        }},
        {{
            "title": "Alternative Industry Position",
            "description": "Role in different industry using transferable skills",
            "required_skills": ["Transferable technical skill", "Industry-specific skill", "Core competency"],
            "matching_skills": ["Relevant transferable skills"],
            "missing_skills": ["Industry-specific skills to learn"],
            "typical_salary": "Salary range in alternative industry",
            "demand_level": "Demand in alternative industry"
        }}
    ],
    "market_insights": {{
        "industry_demand": "Current industry growth rate and future projections",
        "emerging_skills": ["Most in-demand emerging skill", "Future-critical skill"],
        "salary_trends": "Specific salary growth rates and market changes",
        "top_locations": ["Best current location with context", "Emerging hub with growth potential"]
    }}
}}

Analyze this resume information and provide SPECIFIC, DIFFERENTIATED recommendations:
{json.dumps(extracted_info, indent=2)}"""

        guidance_response = model.generate_content(guidance_prompt)
        logger.info(f"AI Guidance Response: {guidance_response.text}")
        
        # Clean the guidance response text
        guidance_text = guidance_response.text.strip()
        if guidance_text.startswith('```json'):
            guidance_text = guidance_text[7:]
        if guidance_text.endswith('```'):
            guidance_text = guidance_text[:-3]
        guidance_text = guidance_text.strip()
        
        try:
            career_guidance = json.loads(guidance_text)
            logger.info("Successfully parsed career guidance")
        except Exception as e:
            logger.error(f"Failed to parse career guidance: {str(e)}")
            career_guidance = {
                "career_paths": [],
                "recommendations": [],
                "job_opportunities": [],
                "market_insights": {
                    "industry_demand": "",
                    "emerging_skills": [],
                    "salary_trends": "",
                    "top_locations": []
                }
            }
        skills = extracted_info.get("skills", [])
        
        # Categorize skills
        technical_skills = [s for s in skills if any(tech in s.lower() for tech in [
            "python", "java", "javascript", "html", "css", "sql", "react", "angular",
            "node", "aws", "azure", "docker", "kubernetes", "git", "api", "programming",
            "development", "software", "web", "database", "cloud", "ai", "ml"
        ])]
        
        soft_skills = [s for s in skills if any(soft in s.lower() for soft in [
            "communication", "leadership", "management", "team", "project", "agile",
            "scrum", "problem solving", "analytical", "organization", "planning"
        ])]
        
        domain_skills = [s for s in skills if s not in technical_skills and s not in soft_skills]
        def calculate_feature_importance():
            features = {
                'technical_skills': len(technical_skills) / 10,  
                'soft_skills': len(soft_skills) / 5,  
                'domain_skills': len(domain_skills) / 5, 
                'experience_quality': 0.0,  
                'education_quality': 0.0, 
            }
            experience = extracted_info.get('experience', '').lower()
            features['experience_quality'] = sum([
                0.2 if 'year' in experience else 0.0,
                0.2 if any(word in experience for word in ['senior', 'lead', 'manager']) else 0.0,
                0.2 if any(word in experience for word in ['project', 'team']) else 0.0,
                0.2 if len(experience.split()) > 30 else 0.0,  
                0.2 if any(word in experience for word in ['achievement', 'success', 'improve', 'develop']) else 0.0
            ])
            education = extracted_info.get('education', '').lower()
            features['education_quality'] = sum([
                0.25 if any(word in education for word in ['bachelor', 'master', 'phd', 'degree']) else 0.0,
                0.25 if any(word in education for word in ['computer', 'engineering', 'science', 'technology']) else 0.0,
                0.25 if len(education.split()) > 20 else 0.0,  
                0.25 if any(word in education for word in ['university', 'college', 'institute']) else 0.0
            ])
            base_score = 0.5  
            contributions = {}
            weights = {
                'technical_skills': 0.3,
                'soft_skills': 0.2,
                'domain_skills': 0.2,
                'experience_quality': 0.2,
                'education_quality': 0.1
            }
            
            total_contribution = 0
            for feature, value in features.items():
                contribution = value * weights[feature]
                contributions[feature] = contribution
                total_contribution += contribution
            normalized_contributions = {k: v/total_contribution for k, v in contributions.items()}
            final_score = base_score + (total_contribution / 2)  
            
            return {
                'score': min(max(final_score, 0), 1),  
                'contributions': normalized_contributions,
                'feature_values': features,
                'explanation': {
                    'technical_skills': f"Found {len(technical_skills)} technical skills",
                    'soft_skills': f"Found {len(soft_skills)} soft skills",
                    'domain_skills': f"Found {len(domain_skills)} domain-specific skills",
                    'experience': "Senior/Lead role detected" if features['experience_quality'] > 0.6 
                                else "Mid-level experience detected" if features['experience_quality'] > 0.3 
                                else "Junior/Entry level experience detected",
                    'education': "Strong educational background" if features['education_quality'] > 0.75
                                else "Standard educational background" if features['education_quality'] > 0.4
                                else "Limited educational background"
                }
            }

        analysis_scores = calculate_feature_importance()
        
        analysis = {
            'extracted_info': extracted_info,
            'career_guidance': career_guidance,
            'score': analysis_scores['score'],
            'feature_importance': analysis_scores['contributions'],
            'feature_values': analysis_scores['feature_values'],
            'explanation': analysis_scores['explanation']
        }

        logger.info(f"Analysis score: {analysis['score']:.2f}")
        logger.info("Feature importance:")
        for feature, importance in analysis['feature_importance'].items():
            logger.info(f"- {feature}: {importance:.2f}")
        logger.info("Explanations:")
        for feature, explanation in analysis['explanation'].items():
            logger.info(f"- {feature}: {explanation}")
            
        return analysis
            
    except Exception as e:
        logger.error(f"Error in AI analysis: {str(e)}")
        return create_default_analysis()

def create_default_analysis():
    """Create a default analysis when AI fails"""
    return {
        "extracted_info": {
            "skills": [],
            "experience": "Could not extract experience information",
            "education": "Could not extract education information"
        },
        "career_guidance": {
            "career_paths": [
                {
                    "role": "Career Analysis Unavailable",
                    "description": "We encountered an error analyzing your resume. Please try again later.",
                    "growth_potential": "Unknown",
                    "salary_range": "Unknown"
                }
            ],
            "recommendations": [
                "Please try uploading your resume again",
                "Make sure your resume is in PDF or DOCX format",
                "Ensure your resume is not password protected"
            ],
            "job_opportunities": [],
            "market_insights": {
                "industry_demand": "Information not available",
                "emerging_skills": [],
                "salary_trends": "Information not available",
                "top_locations": []
            }
        },
        "feature_importance": {
            "Technical Skills": 0.0,
            "Soft Skills": 0.0,
            "Domain Knowledge": 0.0,
            "Experience": 0.0,
            "Education": 0.0
        },
        "score": 0.0
    }

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/analyze', methods=['POST'])
def analyze():
    try:
        if 'resume' not in request.files:
            logger.error("No resume file in request")
            return jsonify({'error': 'No file uploaded'}), 400

        file = request.files['resume']
        if not file or file.filename == '':
            logger.error("Empty filename")
            return jsonify({'error': 'No file selected'}), 400

        filename = secure_filename(file.filename)
        if not allowed_file(filename):
            logger.error(f"Invalid file type: {filename}")
            return jsonify({'error': 'Invalid file type. Please upload PDF or DOCX only.'}), 400

        file_content = file.read()
        if not file_content:
            logger.error("Empty file content")
            return jsonify({'error': 'Empty file uploaded'}), 400
            
        logger.info(f"Processing file: {filename} ({len(file_content)} bytes)")

        file_extension = filename.rsplit('.', 1)[1].lower()
        if file_extension == 'pdf':
            text = extract_text_from_pdf(file_content)
        else:
            text = extract_text_from_docx(file_content)

        if not text:
            logger.error("Text extraction failed")
            return jsonify({
                'error': 'Could not extract text from file. Please ensure:\n' +
                        '1. The file is not corrupted\n' +
                        '2. The file is not password protected\n' +
                        '3. The file contains actual text (not scanned images)\n' +
                        '4. The file is in the correct format (PDF or DOCX)',
                'debug_info': {
                    'filename': filename,
                    'file_size': len(file_content),
                    'file_type': file_extension
                }
            }), 400

        preview = text[:200].replace('\n', ' ')
        logger.info(f"Extracted text preview: {preview}...")

        analysis = analyze_resume_with_ai(text)
        if not analysis:
            logger.error("AI analysis returned None")
            return jsonify({'error': 'Failed to analyze resume content'}), 500

        required_keys = ['score', 'feature_importance', 'extracted_info', 'career_guidance']
        missing_keys = [key for key in required_keys if key not in analysis]
        if missing_keys:
            logger.error(f"Analysis missing required keys: {missing_keys}")
            return jsonify({'error': f'Incomplete analysis data. Missing: {missing_keys}'}), 500

       
        try:
            response = {
                'score': float(analysis['score']),
                'feature_importance': analysis['feature_importance'],
                'extracted_info': analysis['extracted_info'],
                'career_guidance': analysis['career_guidance'],
                'feature_values': analysis['feature_values'],
                'explanation': analysis['explanation'],
                'success': True
            }
            
            
            logger.info(f"Response size: {len(str(response))} characters")
            logger.info(f"Analysis score: {response['score']}")
            logger.info(f"Number of skills: {len(response['extracted_info'].get('skills', []))}")
            
            return jsonify(response)
            
        except Exception as e:
            logger.error(f"Error preparing response: {str(e)}")
            return jsonify({'error': 'Error preparing analysis results'}), 500

    except Exception as e:
        logger.error(f"Error during analysis: {str(e)}")
        return jsonify({'error': str(e)}), 500

if __name__ == '__main__':
    try:
       
        try:
            import pytesseract
            if not os.path.exists("C:\\Program Files\\Tesseract-OCR\\tesseract.exe"):
                logger.warning("Tesseract OCR not found - OCR functionality will be disabled")
            else:
                logger.info("Tesseract OCR found")
                
            # Set Tesseract path
            pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
        except ImportError:
            logger.warning("pytesseract not installed - OCR functionality will be disabled")
            
        # Initialize Gemini
        try:
            genai.configure(api_key=os.getenv('GOOGLE_API_KEY'))
            logger.info("Gemini API initialized successfully")
        except Exception as e:
            logger.error(f"Failed to initialize Gemini API: {str(e)}")
            raise
            
        
        ports = [5000, 8000, 8080, 3000]
        for port in ports:
            try:
                logger.info(f"Attempting to start server on port {port}...")
                app.run(host='127.0.0.1', port=port, debug=True, use_reloader=False)
                break
            except OSError as e:
                if port == ports[-1]: 
                    logger.error(f"Could not bind to any port. Last error: {str(e)}")
                    raise
                else:
                    logger.warning(f"Port {port} is busy, trying next port...")
                    continue
        
    except Exception as e:
        logger.error(f"Failed to start server: {str(e)}")
        raise
